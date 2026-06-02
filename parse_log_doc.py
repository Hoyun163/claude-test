import re
import sys
from docx import Document
from docx.oxml.ns import qn

"""
读取埋点规范文档，将数据转化为一下格式输出：
"美人日志（beauty）" = {"beauty": [['时间', '游戏标识', '游戏版本'], ['yyyy-mm-dd HH:mi:ss', 'appkey', 'version']]}
"""


def is_heading3(xml_element):
    """检查 XML 元素是否为 Heading 3 段落"""
    if xml_element.tag != qn('w:p'):
        return False
    pPr = xml_element.find(qn('w:pPr'))
    if pPr is None:
        return False
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        return False
    return pStyle.get(qn('w:val')) in ('Heading3', '3')


def get_para_text(xml_element):
    """从 w:p XML 元素中提取文本内容"""
    texts = []
    for t in xml_element.iter(qn('w:t')):
        if t.text:
            texts.append(t.text)
    return ''.join(texts)


def extract_english_from_title(title):
    """从标题中提取括号内的英文名称
    例: '美人日志（beauty）' → 'beauty'
         '[C0101]玩家杀怪死亡日志（dead）' → 'dead'
    """
    match = re.search(r'[（(]([a-zA-Z]+)[）)]', title)
    return match.group(1) if match else None


def strip_code_prefix(title):
    """去除标题中的 [CXXXX] 前缀
    例: '[C0101]玩家杀怪死亡日志（dead）' → '玩家杀怪死亡日志（dead）'
    """
    return re.sub(r'^\[C\d+\]', '', title).strip()


def extract_log_name(table):
    """从表格第一行提取日志英文名称
    例: '日志名称：beauty.log.yyyy-mm-dd' → 'beauty'
    """
    row0_text = table.rows[0].cells[0].text
    match = re.search(r'日志名称[：:]\s*(\w+)\.log', row0_text)
    return match.group(1) if match else None


def extract_fields(table):
    """从表格提取'字段名称'和'英文名称'两列数据"""
    field_names = []
    english_names = []
    for row in table.rows[3:]:  # 从第4行（索引3，跳过表头）开始
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 3 and cells[0].isdigit():
            field_names.append(cells[1])
            english_names.append(cells[2])
    return field_names, english_names


def parse_document(docx_path):
    """
    解析文档，遍历 body 子元素建立 Heading 3 与表格的对应关系。
    返回列表，每个元素为:
    {
        'title': str,          # 去除 [CXXXX] 前缀后的标题
        'title_en': str,       # 标题括号中的英文名
        'field_names': list,   # 字段名称列表
        'english_names': list, # 英文名称列表
    }
    """
    doc = Document(docx_path)
    body = doc.element.body

    result = []
    current_title = None
    table_idx = 0

    for child in body:
        tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag_name == 'p':
            if is_heading3(child):
                current_title = get_para_text(child)

        elif tag_name == 'tbl':
            if current_title and table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                log_name = extract_log_name(table)
                title_en = extract_english_from_title(current_title)

                if log_name and title_en and log_name == title_en:
                    field_names, english_names = extract_fields(table)
                    clean_title = strip_code_prefix(current_title)
                    result.append({
                        'title': clean_title,
                        'title_en': title_en,
                        'field_names': field_names,
                        'english_names': english_names,
                    })
            table_idx += 1

    return result


def match_data(all_data, input_names):
    """根据输入名称列表匹配文档数据。
    支持以下输入格式：
    - 完整标题: '美人日志（beauty）', '[C0101]玩家杀怪死亡日志（dead）'
    - 纯英文名: 'beauty', 'dead', 'pkdead'
    """
    matched = []
    for name in input_names:
        name_en = extract_english_from_title(name)
        name_clean = strip_code_prefix(name)

        found = None
        for item in all_data:
            if item['title'] == name_clean or item['title_en'] == name_en:
                found = item
                break

        # 如果前两种方式没匹配到，尝试直接按英文名匹配
        if not found:
            for item in all_data:
                if item['title_en'] == name.strip():
                    found = item
                    break

        if found:
            matched.append(found)
    return matched


def format_output(results):
    """格式化输出为指定格式"""
    lines = []
    for item in results:
        title = item['title']
        title_en = item['title_en']
        field_names = item['field_names']
        english_names = item['english_names']
        lines.append(
            f'"{title}" = {{"{title_en}": [{field_names}, {english_names}]}},'
        )
    return '\n'.join(lines)


def main(input_names, docx_path=r'E:\Claudecode\demo01\规范文档.docx'):   
    # 也可以从命令行参数读取
    # if len(sys.argv) > 1:
    #     input_names = sys.argv[1:]  # python parse_log_doc.py beauty dead shop

    all_data = parse_document(docx_path)
    matched = match_data(all_data, input_names)
    output = format_output(matched)
    print(output)  # 输出位字符串格式


if __name__ == '__main__':
    # ====== 可在此修改变量内容 ======
    # 支持两种格式：
    # 1. 完整标题: "美人日志（beauty）", "[C0101]玩家杀怪死亡日志（dead）"
    # 2. 纯英文名: "beauty", "dead", "pkdead"
    # ===============================
    input_names = [
        '美人日志（beauty）',
        '[C0101]玩家杀怪死亡日志（dead）',
        '[C0102]玩家pk死亡日志（pkdead）',
        '[C0103]元宝商城日志（shop）',
        '[C0104]寄售日志（consignment）',
        '[C0105]放弃任务日志（canceltask）',
        '[C0106]伙伴日志（fellow）',
        '[C0107]玩家匹配（帮会)日志（match）',
    ]
    main(input_names)
